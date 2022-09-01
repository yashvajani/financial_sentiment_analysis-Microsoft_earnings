# %% Creating a function to extract and store textual data

'''
The function below extracts textual data from microsoft earnings reports in the 
specified working directory and stores it in a dataframe with corresponding information 
about the sentence position in the text
'''
# importing libraries
import glob, os
import re
import pandas as pd
import spacy
from docx import Document
nlp = spacy.load('en_core_web_sm')

class ExtractData():
    def extractText(self, path, quarter="PASS"):
        '''
        path = specified working directory
        quarter = relevant quarter of Microsoft earnings release

        To extract data relating to Microsoft earnings, add the "PressReleaseFYXXQX" to a folder and then specify it as the working directory.
        This function extracts text data from Word files in the specified working directory on a sentence basis
        '''
        wk_dir = path
        os.chdir(wk_dir)
        in_fs = glob.glob('**/*.docx', recursive=True)

        # Creating output dataframe
        out_df = pd.DataFrame()
        read_text = {"quarter":[], "paragraph":[], "sentence":[], "text":[]}

        for file in sorted(in_fs):
            folderpath = wk_dir
            in_f = os.path.join(folderpath,file)
            f = open(in_f, 'rb')
            doc = Document(f)
            f.close()

            paragraph_count = 1
            for para in doc.paragraphs:
                nlp_var = nlp(para.text)
                sentence_count = 1
                for sentence in nlp_var.sents:
                    read_text['quarter'].append(file[12:-5]) # recording quarter corresponding to data-point
                    read_text['paragraph'].append(paragraph_count) # recording the paragraph count corresponding to data-point
                    read_text['sentence'].append(sentence_count) # recording the sentence count corresponding to data-point
                    read_text['text'].append(sentence.text) # recording the text
                    sentence_count += 1
                paragraph_count += 1
        
        df = pd.DataFrame(read_text)
        out_df = pd.concat([out_df, df])
        if quarter!="PASS":
            out_df = out_df[out_df['quarter'] == quarter]
        return out_df
    
    def extractParagraph(self, path, quarter="PASS"):
        '''
        path = specified working directory
        quarter = relevant quarter of Microsoft earnings release

        To extract data relating to Microsoft earnings, add the "PressReleaseFYXXQX" to a folder and then specify it as the working directory.
        This function extracts text data from Word files in the specified working directory on a paragraph basis
        '''
        wk_dir = path
        os.chdir(wk_dir)
        in_fs = glob.glob('**/*.docx', recursive=True)

        out_df = pd.DataFrame()
        read_text = {"quarter":[], "paragraph":[], "text":[]}

        for file in sorted(in_fs):
            folderpath = wk_dir
            in_f = os.path.join(folderpath,file)
            f = open(in_f, 'rb')
            doc = Document(f)
            f.close()

            paragraph_count = 1
            for para in doc.paragraphs:
                read_text['quarter'].append(file[12:-5]) # recording quarter corresponding to data-point
                read_text['paragraph'].append(paragraph_count) # recording the paragraph count corresponding to data-point
                read_text['text'].append(para.text)  # recording the text
                paragraph_count += 1
            
        df = pd.DataFrame(read_text)
        out_df = pd.concat([out_df, df])
        if quarter!="PASS":
            out_df = out_df[out_df['quarter'] == quarter]
        return out_df
    
    def extractFinancials(self, path, quarter="PASS", diluted = True):
        '''
        path = specified working directory
        quarter = relevant quarter of Microsoft earnings release
        diluted = 'True' to extract diluted EPS and 'False' to extract basic EPS

        To extract data relating to Microsoft earnings, add the "PressReleaseFYXXQX" to a folder and then specify it as the working directory.
        This function extracts table data from Word files in the specified working directory on a row-by-row basis
        '''
        wk_dir = path
        os.chdir(wk_dir)
        in_fs = glob.glob('**/*.docx', recursive=True)

        out_df = pd.DataFrame()
        read_text = {"quarter":[], "table":[], "row":[], "data":[]}

        for file in sorted(in_fs):
            folderpath = wk_dir
            in_f = os.path.join(folderpath, file)
            f = open(in_f, 'rb')
            doc = Document(f)
            f.close()
            
            table_count = 1
            for table in doc.tables:
                row_count = 1
                for row in table.rows:
                    row_data = []
                    read_text['quarter'].append(file[12:-5]) # recording quarter corresponding to data-point
                    read_text['table'].append(table_count) # recording the table count corresponding to data-point
                    read_text['row'].append(row_count) # recording the row count corresponding to data-point
                    for cell in row.cells:
                        row_data.append(cell.text) # creating a list of row data
                    row_count += 1
                    read_text['data'].append(row_data) # recording row data
                table_count += 1

        df = pd.DataFrame(read_text)
        out_df = pd.concat([out_df, df])
        if quarter!="PASS":
            out_df = out_df[out_df['quarter'] == quarter]

        '''
        As this function is meant to extract data from Microsoft's earnings release documents, the proceding code extracts 
        the EPS for the quarters under consideration
        '''
        temp_df = pd.DataFrame()
        eps_tags = ['Earnings per share:', 'Earnings (loss) per share:']
        temp_df["eps_tag"] = out_df['data'].apply(lambda x: 1 if any(i in x for i in eps_tags) else 0)
        eps_idx = (temp_df.index[temp_df['eps_tag'] == 1].tolist())
        eps=[]
        if diluted == False:
            for index in eps_idx:
                idx = index+1
                eps.append(out_df['data'][idx][1])
        else:
            for index in eps_idx:
                idx = index+2
                eps.append(out_df['data'][idx][1])
        
        def findBrackets(st): # a function that returns -1 for figures with brackets and 1 otherwise
            for i in range(len(st)):
                if st[i] == '(':
                    return -1
            return 1

        eps_list = pd.Series(eps)
        numbers = []
        for i in eps_list:
            s = [float(s) * findBrackets(i) for s in re.findall(r'-?\d+\.?\d*', i)] 
            numbers.append(s)
        numbers = [i for list in numbers for i in list]
        
        earnings = pd.DataFrame()
        earnings['quarter'] = [qt for qt in out_df['quarter'].unique()]
        earnings['eps'] = numbers
        
        del read_text, temp_df, row_data, df, eps, eps_list, numbers
        return out_df, earnings